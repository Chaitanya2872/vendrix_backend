from datetime import datetime, timezone
from pathlib import Path
import shutil
from io import BytesIO
from uuid import uuid4
from typing import Any
from fastapi import APIRouter, BackgroundTasks, Depends, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import select
from app.common.dependencies import current_user
from app.core.config import settings
from app.db.session import get_db
from app.models import AuditLog, Document, User

router = APIRouter(prefix="/documents", tags=["documents"])
class DocumentReview(BaseModel): fields: dict[str, Any]

@router.get("")
def list_documents(db: Session = Depends(get_db), _: User = Depends(current_user)):
    return db.scalars(select(Document).order_by(Document.created_at.desc())).all()

@router.post("", status_code=201)
def upload(document_type: str, background_tasks: BackgroundTasks, file: UploadFile = File(...), db: Session = Depends(get_db), user: User = Depends(current_user)):
    extension = Path(file.filename or "").suffix.lower()
    supported_extensions = {".pdf", ".jpg", ".jpeg", ".png", ".webp", ".docx", ".xlsx"}
    if extension not in supported_extensions:
        raise HTTPException(415, "Supported files: PDF, JPG, PNG, WEBP, DOCX, XLSX")
    key = f"{user.id}/{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}_{file.filename}"; path = Path(settings.storage_path) / key; path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("wb") as output: shutil.copyfileobj(file.file, output)
    document = Document(filename=file.filename, object_key=key, content_type=file.content_type, document_type=document_type.upper(), owner_id=user.id)
    db.add(document); db.flush(); db.add(AuditLog(actor_id=user.id, action="UPLOAD", resource_type="documents", resource_id=document.id)); db.commit(); db.refresh(document)
    from app.workers.document_tasks import process_vendor_document, process_vendor_document_now
    if settings.celery_enabled:
        try:
            process_vendor_document.delay(document.id)
        except Exception:
            # Keep the request fast if the queue is temporarily unavailable.
            background_tasks.add_task(process_vendor_document_now, document.id)
    else:
        # Local development: send the response first, then process in-process.
        background_tasks.add_task(process_vendor_document_now, document.id)
    return document

@router.get("/{document_id}")
def get_document(document_id: str, db: Session = Depends(get_db), _: User = Depends(current_user)):
    document = db.get(Document, document_id)
    if not document: raise HTTPException(404, "Document not found")
    return document

@router.get("/{document_id}/download")
def download_document(document_id: str, db: Session = Depends(get_db), _: User = Depends(current_user)):
    document = db.get(Document, document_id)
    if not document: raise HTTPException(404, "Document not found")
    stored_file = Path(settings.storage_path) / document.object_key
    if not stored_file.exists(): raise HTTPException(404, "Stored file not found")
    return FileResponse(stored_file, media_type=document.content_type or "application/octet-stream", filename=document.filename, content_disposition_type="inline")

@router.get("/{document_id}/preview")
def document_preview(document_id: str, db: Session = Depends(get_db), _: User = Depends(current_user)):
    document = db.get(Document, document_id)
    if not document: raise HTTPException(404, "Document not found")
    stored_file = Path(settings.storage_path) / document.object_key
    if not stored_file.exists(): raise HTTPException(404, "Stored file not found")
    extension = stored_file.suffix.lower()
    if extension == ".docx":
        from docx import Document as WordDocument
        lines = [paragraph.text for paragraph in WordDocument(stored_file).paragraphs if paragraph.text.strip()]
        return {"type": "document", "lines": lines}
    if extension == ".xlsx":
        from openpyxl import load_workbook
        workbook = load_workbook(BytesIO(stored_file.read_bytes()), read_only=True, data_only=True)
        sheet = workbook.active
        rows = [["" if value is None else str(value) for value in row] for row in sheet.iter_rows(values_only=True, max_row=100, max_col=20)]
        return {"type": "spreadsheet", "sheet": sheet.title, "rows": rows}
    return {"type": "binary"}

@router.delete("/{document_id}", status_code=204)
def delete_document(document_id: str, db: Session = Depends(get_db), user: User = Depends(current_user)):
    document = db.get(Document, document_id)
    if not document: raise HTTPException(404, "Document not found")
    stored_file = Path(settings.storage_path) / document.object_key
    if stored_file.exists(): stored_file.unlink()
    db.add(AuditLog(actor_id=user.id, action="DELETE", resource_type="documents", resource_id=document.id)); db.delete(document); db.commit()

@router.post("/{document_id}/review")
def review(document_id: str, body: DocumentReview, db: Session = Depends(get_db), user: User = Depends(current_user)):
    document = db.get(Document, document_id)
    if not document: raise HTTPException(404, "Document not found")
    document.extracted_fields, document.status, document.review_confirmed_at = body.fields, "CONFIRMED", datetime.now(timezone.utc)
    db.add(AuditLog(actor_id=user.id, action="CONFIRM_EXTRACTION", resource_type="documents", resource_id=document.id)); db.commit(); db.refresh(document); return document
